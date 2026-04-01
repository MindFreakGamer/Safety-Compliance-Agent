"""
Document Parser Module
Extracts text from safety documents (PDF, DOCX, TXT) and splits into logical sections.

Design Decision: Text-only extraction for v1. Vision-based parsing (sending PDF pages
as images to Claude) would handle diagrams and flowcharts — planned for v2.
"""

import os
import re
from dataclasses import dataclass, field
from typing import Optional


@dataclass
class DocumentSection:
    """Represents a logical section of the safety document."""
    section_number: str
    title: str
    content: str
    page_number: Optional[int] = None

    def __str__(self):
        return f"[Section {self.section_number}] {self.title}\n{self.content}"


@dataclass
class ParsedDocument:
    """Represents a fully parsed safety document."""
    filename: str
    file_type: str
    full_text: str
    sections: list = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

    @property
    def section_count(self):
        return len(self.sections)

    @property
    def word_count(self):
        return len(self.full_text.split())


def parse_document(file_path: str) -> ParsedDocument:
    """
    Parse a safety document and extract text content.
    
    Supports: .txt, .pdf, .docx
    Returns a ParsedDocument with full text and split sections.
    """
    ext = os.path.splitext(file_path)[1].lower()
    filename = os.path.basename(file_path)

    if ext == ".txt":
        full_text = _parse_txt(file_path)
    elif ext == ".pdf":
        full_text = _parse_pdf(file_path)
    elif ext == ".docx":
        full_text = _parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .txt, .pdf, .docx")

    sections = _split_into_sections(full_text)
    metadata = _extract_metadata(full_text)

    return ParsedDocument(
        filename=filename,
        file_type=ext,
        full_text=full_text,
        sections=sections,
        metadata=metadata,
    )


def parse_text_directly(text: str, source_name: str = "direct_input") -> ParsedDocument:
    """Parse text content directly (for pasted text in the UI)."""
    sections = _split_into_sections(text)
    metadata = _extract_metadata(text)

    return ParsedDocument(
        filename=source_name,
        file_type=".txt",
        full_text=text,
        sections=sections,
        metadata=metadata,
    )


def _parse_txt(file_path: str) -> str:
    """Extract text from a plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError(
            "pdfplumber is required for PDF parsing. Install with: pip install pdfplumber"
        )

    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {i + 1} ---\n{page_text}")

    if not text_parts:
        raise ValueError(f"Could not extract any text from PDF: {file_path}")

    return "\n\n".join(text_parts)


def _parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX parsing. Install with: pip install python-docx"
        )

    doc = Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    if not paragraphs:
        raise ValueError(f"Could not extract any text from DOCX: {file_path}")

    return "\n\n".join(paragraphs)


def _split_into_sections(text: str) -> list:
    """
    Split document text into logical sections.
    
    Handles common SOP formatting patterns:
    - Numbered sections (1. TITLE, 2.1 Subtitle)
    - Section headers with separators (=== or ---)
    - All-caps headers
    """
    sections = []

    # Pattern: Match numbered sections like "1.", "4.1", "5.2.3" followed by title text
    # Also handles lines with separator bars (===, ---) before/after headings
    section_pattern = re.compile(
        r"(?:^[=\-]{3,}\s*\n)?"  # Optional separator line before
        r"^(\d+(?:\.\d+)*)\s*[.\):]?\s+"  # Section number
        r"([A-Z][^\n]+)"  # Title (starts with capital letter)
        r"(?:\n[=\-]{3,})?"  # Optional separator line after
        ,
        re.MULTILINE,
    )

    matches = list(section_pattern.finditer(text))

    if not matches:
        # Fallback: treat entire document as one section
        sections.append(
            DocumentSection(
                section_number="1",
                title="Full Document",
                content=text.strip(),
            )
        )
        return sections

    for i, match in enumerate(matches):
        section_num = match.group(1)
        title = match.group(2).strip().rstrip("=").rstrip("-").strip()

        # Content runs from end of this match to start of next match (or end of text)
        content_start = match.end()
        content_end = matches[i + 1].start() if i + 1 < len(matches) else len(text)

        content = text[content_start:content_end].strip()
        # Clean up separator lines from content
        content = re.sub(r"^[=\-]{3,}\s*$", "", content, flags=re.MULTILINE).strip()

        sections.append(
            DocumentSection(
                section_number=section_num,
                title=title,
                content=content,
            )
        )

    return sections


def _extract_metadata(text: str) -> dict:
    """Extract document metadata like title, revision, date from the header."""
    metadata = {}

    # Try to extract document title (usually first non-empty line)
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if lines:
        # Skip separator lines
        for line in lines:
            if not re.match(r"^[=\-]{3,}$", line):
                metadata["title"] = line
                break

    # Try to extract common metadata fields
    patterns = {
        "document_number": r"Document\s*(?:No|Number|#)[:\s]*([^\n]+)",
        "revision": r"Revision[:\s]*([^\n]+)",
        "effective_date": r"Effective\s*Date[:\s]*([^\n]+)",
        "prepared_by": r"Prepared\s*by[:\s]*([^\n]+)",
        "approved_by": r"Approved\s*by[:\s]*([^\n]+)",
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            metadata[key] = match.group(1).strip()

    return metadata
